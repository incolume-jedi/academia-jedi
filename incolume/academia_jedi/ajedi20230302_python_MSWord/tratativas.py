"""Gerar documentos DOCX com python-docx.


https://python-docx.readthedocs.io/en/latest/.
"""

import datetime as dt
import inspect
import logging
import typing
from pathlib import Path

import pytz
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

# ruff: noqa: A001 A002 ANN001 ANN002 ANN003 ANN201 ANN202 ANN204 ANN401 ARG001 ARG002 ASYNC101 B007 B008 B009 B011 B015 B904 B905 BLE001 C408 C419 C901 D100 D101 D102 D103 D104 D105 D107 D205 D402 D415 D419 DTZ001 DTZ003 DTZ005 DTZ007 E501 E741 EM101 EM102 ERA001 EXE005 F402 F403 F405 F601 F811 F821 F841 FBT001 FBT002 FBT003 FIX002 G001 G002 G004 N801 N802 N805 N806 N816 N999 NPY002 PD901 PERF203 PERF401 PERF402 PIE796 PLE1205 PLR0913 PLR1714 PLR2004 PLW0602 PLW0603 PLW2901 PT004 PT006 PT012 PT015 PTH118 PTH123 PYI024 PYI041 RET503 RET504 RUF001 RUF012 RUF013 S101 S113 S201 S301 S307 S310 S311 S602 S603 S605 S607 S608 SIM103 SIM109 SIM113 SIM115 SIM117 SLF001 SLOT000 T201 T203 TCH003 TD002 TD003 TD004 TRY002 TRY003 TRY300 TRY301 TRY401 W293
texto: str = Path('texto_template.txt').read_text()
bottom: str = (
    'Generated 15 paragraphs, ' '1394 words, 9332 bytes of Lorem Ipsum at {}'
)


def tratativa1():
    """Criação de documento MSWord."""
    documento = Document()
    timestamp = dt.datetime.now(tz=pytz.timezone('America/Sao_Paulo'))
    bottom.format(timestamp.isoformat())
    documento.add_paragraph(texto)
    documento.add_paragraph(bottom)
    documento.save(Path(__file__).parent / f'{inspect.stack()[0][3]}.docx')


def tratativa2():
    """Estilo no DOCX."""
    documento = Document()
    texto = """Olá Mundo!!!"""
    paragraf = documento.add_paragraph(texto)
    paragraf.style = documento.styles.add_style(
        'Mystyle',
        WD_STYLE_TYPE.PARAGRAPH,
    )
    paragraf.style.font.name = 'Algerian'
    paragraf.style.font.size = Pt(15)
    paragraf.style.font.bold = True
    paragraf.style.font.italic = True

    documento.save(Path(__file__).parent / f'{inspect.stack()[0][3]}.docx')


def tratativa3():
    """Estilo no DOCX."""
    documento = Document()
    texto = """Olá Mundo!!!"""
    paragraf = documento.add_paragraph(texto)
    paragraf.style = documento.styles.add_style(
        (mystyle := inspect.stack()[0][3]),
        WD_STYLE_TYPE.PARAGRAPH,
    )
    paragraf.style.font.name = 'Algerian'
    paragraf.style.font.size = Pt(15)
    paragraf.style.font.bold = True
    paragraf.style.font.italic = True
    paragraf.style.font.underline = True
    paragraf.style.font.color.rgb = RGBColor(255, 0, 0)

    # sem estilo aplicado
    documento.add_paragraph('Novo paragrafo.')

    # com estilo previamente criado
    documento.add_paragraph('Novo paragrafo.', mystyle)

    documento.save(Path(__file__).parent / f'{mystyle}.docx')


def tratativa4():
    """Exibir estilos disponiveis."""
    documento = Document()
    for estilo in documento.styles:
        print(estilo)


def tratativa5():
    """Utilizar estilos pré-existentes."""
    documento = Document()

    documento.add_paragraph('Cabeçalho', 'Header')
    documento.add_paragraph('Primeiro paragrafo.', 'Heading 1')
    documento.add_paragraph('Segundo paragrafo.')
    documento.add_paragraph('Terceiro paragrafo.')
    documento.add_paragraph('Quarto paragrafo.')

    documento.save(Path(__file__).parent / f'{inspect.stack()[0][3]}.docx')


def tratativa6():
    """Criar um documento Word.

    Este exemplo cria um arquivo a partir de uma configuração
    de estilo preexistente em outro arquivo.
    """
    documento = Document(Path(__file__).parent / 'tratativa2.docx')
    documento.add_paragraph('Cabeçalho', 'Heading 1')
    documento.add_paragraph('Primeiro paragrafo.', 'Mystyle')
    documento.add_paragraph('Segundo paragrafo.', 'Mystyle')
    documento.add_paragraph('Terceiro paragrafo.', 'Mystyle')
    documento.add_paragraph('Quarto paragrafo.', 'Mystyle')

    documento.save(Path(__file__).parent / f'{inspect.stack()[0][3]}.docx')


def tratativa7():
    """Formatação Negrito em texto."""
    documento = Document()
    paragraf = documento.add_paragraph('Texto gerado em ')
    timestamp = dt.datetime.now(tz=pytz.timezone('America/Sao_Paulo'))
    paragraf.add_run(f'{timestamp.isoformat()} ').bold = True
    paragraf.add_run('através do ')
    paragraf.add_run('Python').italic = True

    documento.save(Path(__file__).parent / f'{inspect.stack()[0][3]}.docx')


def tratativa8():
    """Quebra de página."""
    documento = Document()

    documento.add_paragraph('Título', 'Heading 1')

    paragraf = documento.add_paragraph('Texto gerado em ')
    timestamp = dt.datetime.now(tz=pytz.timezone('America/Sao_Paulo'))
    paragraf.add_run(f'{timestamp.isoformat()}').bold = True
    paragraf.add_run('através do ')
    paragraf.add_run('Python').italic = True

    documento.add_page_break()
    documento.add_paragraph('Título', 'Heading 1')
    documento.add_paragraph('Primeiro paragrafo.')
    documento.add_paragraph('Segundo paragrafo.')
    documento.add_paragraph('Terceiro paragrafo.')
    documento.add_paragraph('Quarto paragrafo.')

    documento.save(Path(__file__).parent / f'{inspect.stack()[0][3]}.docx')


def tratativa9():
    """Controle de margem e seções de página."""
    documento = Document()
    for section in documento.sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(0.5)

    timestamp = dt.datetime.now(tz=pytz.timezone('America/Sao_Paulo'))
    paragraf = documento.add_paragraph('Texto gerado em ')
    paragraf.add_run(
        f'{timestamp.isoformat()} ',
    ).bold = True
    paragraf.add_run('através do ')
    paragraf.add_run('Python').italic = True

    documento.save(Path(__file__).parent / f'{inspect.stack()[0][3]}.docx')


def tratativa10():
    """Controle de margem e seções de página."""
    documento = Document()
    for section in documento.sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(0.5)

    timestamp = dt.datetime.now(tz=pytz.timezone('America/Sao_Paulo'))
    paragraf = documento.add_paragraph(texto)
    paragraf.add_run('\nTexto gerado em ')
    paragraf.add_run(
        f'{timestamp.isoformat()} ',
    ).bold = True
    paragraf.add_run('através do ')
    paragraf.add_run('Python').italic = True
    paragraf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    documento.save(Path(__file__).parent / f'{inspect.stack()[0][3]}.docx')


def tratativa11():
    """Imagem."""
    documento = Document()
    content = texto.split('\n')
    logging.debug(content)
    documento.add_paragraph(content.pop(0), 'Heading 1')
    documento.add_paragraph(content.pop(0))
    documento.add_paragraph(content.pop(0))

    documento.add_picture(
        Path(__file__).parent.joinpath('imagem.png').as_posix(),
    )
    documento.add_paragraph(content.pop(0))

    documento.save(Path(__file__).parent / f'{inspect.stack()[0][3]}.docx')


def tratativa12():
    """Imagem."""
    documento = Document()
    content = texto.split('\n')
    documento.add_paragraph(content.pop(0), 'Heading 1')
    documento.add_paragraph(content.pop(0))
    documento.add_paragraph(content.pop(0))

    documento.add_picture(
        Path(__file__).parent.joinpath('imagem.png').as_posix(),
        width=Cm(4),
        height=Cm(4),
    )
    documento.add_paragraph(content.pop(0))

    documento.save(Path(__file__).parent / f'{inspect.stack()[0][3]}.docx')


def tratativa13():
    """Imagem."""
    documento = Document()
    content = texto.split('\n')
    documento.add_paragraph(content.pop(0), 'Heading 1')
    documento.add_paragraph(content.pop(0))
    documento.add_paragraph(content.pop(0))

    image = documento.add_picture(
        Path(__file__).parent.joinpath('imagem.png').as_posix(),
        width=Cm(4),
        height=Cm(4),
    )
    image.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for cont in content:
        parag = documento.add_paragraph(cont)
        parag.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    documento.save(Path(__file__).parent / f'{inspect.stack()[0][3]}.docx')


def tratativa14():
    """Tabela."""
    documento = Document()
    paragrafos = texto.format(
        dt.datetime.now(tz=pytz.timezone('America/Sao_Paulo')).isoformat(),
    ).split('\n')

    documento.add_paragraph(paragrafos.pop(0), 'Heading 1')

    for parag in paragrafos[:2]:
        paragraf = documento.add_paragraph(parag)
        paragraf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragrafos.pop(0)

    image = documento.add_picture(
        Path(__file__).parent.joinpath('imagem.png').as_posix(),
        width=Cm(4),
        height=Cm(4),
    )
    image.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cont in paragrafos[:5]:
        parag = documento.add_paragraph(cont)
        parag.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragrafos.pop(0)
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam'),
    )

    table = documento.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, _id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = _id
        row_cells[2].text = desc

    for cont in paragrafos:
        parag = documento.add_paragraph(cont)
        parag.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    documento.save(Path(__file__).parent / f'{inspect.stack()[0][3]}.docx')


def tratativa15():
    """Tabela: Estilos disponíveis."""
    documento = Document()
    for estilo in (x for x in documento.styles):
        print(estilo)


def tratativa16():
    """Tabela."""
    documento = Document()
    paragrafos = texto.format(
        dt.datetime.now(tz=pytz.timezone('America/Sao_Paulo')).isoformat(),
    ).split('\n')

    documento.add_paragraph(paragrafos.pop(0), 'Heading 1')

    for parag in paragrafos[:2]:
        paragraf = documento.add_paragraph(parag)
        paragraf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragrafos.pop(0)

    image = documento.add_picture(
        Path(__file__).parent.joinpath('imagem.png').as_posix(),
        width=Cm(4),
        height=Cm(4),
    )
    image.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for cont in paragrafos[:5]:
        parag = documento.add_paragraph(cont)
        parag.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragrafos.pop(0)
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam'),
    )

    table = documento.add_table(rows=1, cols=3, style='Light List Accent 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, _id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = _id
        row_cells[2].text = desc

    for cont in paragrafos:
        parag = documento.add_paragraph(cont)
        parag.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    documento.save(Path(__file__).parent / f'{inspect.stack()[0][3]}.docx')


def tratativa17():
    """Mala direta MSWord via Python."""
    contrato = Document(Path(__file__).parent / 'Contrato.docx')

    nome = 'Lira da Hashtag'
    item1 = 'Serviço de Treinamento em Excel'
    item2 = 'Apostila Completa de Excel'
    item3 = 'Serviço de Treinamentos de Python'

    timestamp = dt.datetime.now(tz=pytz.timezone('America/Sao_Paulo'))
    macros = {
        'XXXX': nome,
        'YYYY': item1,
        'ZZZZ': item2,
        'WWWW': item3,
        'DD': str(timestamp.day),
        'MM': str(timestamp.month),
        'AAAA': str(timestamp.year),
    }
    for paragraf in contrato.paragraphs:
        print(paragraf)
        if 'XXXX' in paragraf.text:
            paragraf.text = paragraf.text.replace('XXXX', macros['XXXX'])

    contrato.save(Path(__file__).parent / f'{inspect.stack()[0][3]}.docx')


def tratativa18():
    """Mala direta MSWord via Python."""
    contrato = Document('Contrato.docx')

    nome = 'Lira da Hashtag'
    item1 = 'Serviço de Treinamento em Excel'
    item2 = 'Apostila Completa de Excel'
    item3 = 'Serviço de Treinamentos de Python'

    timestamp = dt.datetime.now(tz=pytz.timezone('America/Sao_Paulo'))
    dicionario_valores = {
        'XXXX': nome,
        'YYYY': item1,
        'ZZZZ': item2,
        'WWWW': item3,
        'DD': str(timestamp.day),
        'MM': str(timestamp.month),
        'AAAA': str(timestamp.year),
    }
    for paragraf in contrato.paragraphs:
        for key in dicionario_valores:
            if key in paragraf.text:
                logging.debug(paragraf, key)
                paragraf.text = paragraf.text.replace(
                    key,
                    dicionario_valores.get(key),
                )

    contrato.save(Path(__file__).parent / f'{inspect.stack()[0][3]}.docx')


def run():
    """Running it."""
    functions: list[typing.Callable] = [
        value
        for key, value in globals().items()
        if key.__contains__('tratativa')
    ]
    for func in functions:
        logging.debug('%s: %s', type(func), func.__name__)
        print(f'--- {func.__name__} ---')
        print(f'    >>> {func.__doc__}')
        try:
            if result := func():
                print(result)
        except (TypeError, ValueError) as e:
            logging.exception('%s', e.__class__.__name__)
        finally:
            logging.debug('%s finalizada.', func.__name__)
        print('------\n')


if __name__ == '__main__':  # pragma: no cover
    run()
