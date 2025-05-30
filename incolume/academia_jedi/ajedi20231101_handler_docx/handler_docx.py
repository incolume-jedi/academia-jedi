import logging
import sys
from os import getenv
from pathlib import Path

from dotenv import load_dotenv
from icecream import ic

try:
    from docx import Document
    from docx.shared import Inches
except (ImportError, ModuleNotFoundError):
    msg = 'ModuleNotFoundError(docx): No package named `python-docx`'
    logging.debug(ic(msg))
    sys.exit(msg)


load_dotenv()

# ruff: noqa: A001, ANN201, D100, ERA001, RUF013


def example_default(fout: Path = None, fimg: Path = None):
    """Exemplo oficial da documentação python-docx."""
    fout = Path(getenv('HANDLER_DOCX_FILENAME'))
    fimg = fimg or Path(__file__).parent / 'img' / 'image001.png'

    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list',
        style='List Bullet',
    )
    document.add_paragraph(
        'first item in ordered list',
        style='List Number',
    )

    document.add_picture(fimg.as_posix(), width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam'),
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save(fout.as_posix())


def docx_basic(fout: Path = None, content: list = None) -> bool:
    """Criar um docx básico."""
    fout = fout or Path(getenv('HANDLER_DOCX_FILENAME'))

    document = Document()

    document.add_heading(content.pop(0))

    p = document.add_paragraph('')
    p.add_run(content.pop(0)).italic = True

    for conteudo in content:
        document.add_paragraph(conteudo)

    document.save(fout.as_posix())

    return fout.is_file()


def run():
    """Run it."""
    # example_default()

    file = Path(__file__).parent.joinpath('texto.docx')
    with Path(__file__).parent.joinpath('texto.txt').open() as f:
        docx_basic(file, [x.strip() for x in f.readlines()])


if __name__ == '__main__':
    run()
